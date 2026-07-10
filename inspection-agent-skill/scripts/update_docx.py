#!/usr/bin/env python
"""
update_docx.py — 更新巡查台账 Word 文档

将巡检截图贴入 Word 台账对应摄像头名称行的"监控截图"列。

文档结构（以品牌配置中的模版为例）：
  - 第 0 行：巡查日期（跨列合并）
  - 第 1 行：表头（序号 | 企业名称 | 摄像头名称 | 是否超阈值 | 其他异常情况 | 监控截图 | 处置情况）
  - 第 2 行起：数据行
  - 列索引 2 = 摄像头名称
  - 列索引 5 = 监控截图（插入图片位置）

用法：
  python scripts/update_docx.py \
    --docx <台账模版路径> \
    --results ./screenshots/inspection_20260629.json \
    --output <输出路径>

  python scripts/update_docx.py \
    --docx <台账模版路径> \
    --results-list '{"摄像头名称": "/path/to/img.jpg", ...}' \
    --output <输出路径>
"""
import argparse
import json
import os
import re
import sys
from datetime import date

# ─── 依赖检查 ─────────────────────────────────────────────────────

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: 请安装 python-docx: pip install python-docx")
    sys.exit(1)

# ─── 工具函数 ─────────────────────────────────────────────────────

def _normalize(name: str) -> str:
    """规范化摄像头名称用于模糊匹配：去除空白、统一连字符"""
    if not name:
        return ""
    name = name.strip()
    # 全角转半角
    name = name.replace("（", "(").replace("）", ")").replace("－", "-").replace("—", "-")
    # 压缩多余空格
    name = re.sub(r"\s+", "", name)
    return name.lower()


def _find_row_by_camera_name(table, target_name: str) -> int:
    """
    在表格第 3 列（索引 2）中查找目标摄像头名称所在行号。
    优先精确匹配，其次规范化模糊匹配。
    返回行索引（从 0 起），未找到返回 -1。
    """
    target_norm = _normalize(target_name)

    # 第一遍：精确匹配
    for ri, row in enumerate(table.rows):
        cell_text = row.cells[2].text.strip()
        if cell_text == target_name:
            return ri

    # 第二遍：规范化匹配
    for ri, row in enumerate(table.rows):
        cell_text = row.cells[2].text.strip()
        if _normalize(cell_text) == target_norm:
            return ri

    return -1


def _insert_image_to_cell(cell, image_path: str, max_width_inches: float = 1.8):
    """
    向单元格插入图片。
    - 清空单元格现有内容
    - 插入图片，宽度限制在 max_width_inches 英寸内
    - 图片居中对齐
    """
    # 清空段落文字
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""

    # 取第一个段落用于插入图片，设置居中
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 清除该段落中已有的图片（inline drawing）
    p_elem = para._p
    for r_elem in list(p_elem):
        if r_elem.tag.endswith("}r"):
            has_drawing = r_elem.find(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
            )
            if has_drawing is not None:
                p_elem.remove(r_elem)

    # 插入图片
    run = para.add_run()
    run.add_picture(image_path, width=Inches(max_width_inches))


def _set_date_row(table, today_str: str):
    """
    更新第 0 行（日期行）的日期文本，格式为 "巡查日期：YYYY.M.D"
    """
    try:
        date_cell = table.rows[0].cells[0]
        for para in date_cell.paragraphs:
            full_text = para.text
            if "巡查日期" in full_text:
                for run in para.runs:
                    if "巡查日期" in run.text or run.text.strip():
                        run.text = f"巡查日期：{today_str}"
                        break
                break
    except Exception:
        pass  # 日期行更新失败不影响主流程


# ─── 主函数 ──────────────────────────────────────────────────────

def update_docx(
    docx_path: str,
    results: list,
    output_path: str,
    today_str: str = None,
    img_width_inches: float = 1.8,
    skip_offline: bool = True,
) -> dict:
    """
    将巡检截图贴入 Word 台账。

    参数：
      docx_path      — 输入 Word 文档路径（模版）
      results        — 巡检结果列表，每项包含 cameraName / screenshotPath / status
      output_path    — 输出 Word 文档路径
      today_str      — 巡查日期字符串，如 "2026.6.29"（默认今天）
      img_width_inches — 图片宽度（英寸），默认 1.8
      skip_offline   — 是否跳过离线/无截图的摄像头（默认 True）

    返回 dict：
      matched    — 成功插入图片的行数
      unmatched  — 在文档中未找到对应行的摄像头名称列表
      skipped    — 跳过（无截图）的数量
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Word 文档不存在: {docx_path}")

    doc = Document(docx_path)
    table = doc.tables[0]

    # 更新日期行
    if today_str is None:
        d = date.today()
        today_str = f"{d.year}.{d.month}.{d.day}"
    _set_date_row(table, today_str)

    matched = 0
    unmatched = []
    skipped = 0

    for result in results:
        cam_name = result.get("cameraName", "").strip()
        img_path = result.get("screenshotPath") or ""
        status = result.get("status", "offline")

        if not cam_name:
            continue

        # 无截图时跳过
        if not img_path or not os.path.exists(img_path):
            if skip_offline:
                skipped += 1
                continue
            else:
                skipped += 1
                continue

        row_idx = _find_row_by_camera_name(table, cam_name)
        if row_idx == -1:
            unmatched.append(cam_name)
            print(f"  [WARN] 未找到匹配行: {cam_name!r}")
            continue

        # 插入截图到"监控截图"列（索引 5）
        cell = table.rows[row_idx].cells[5]
        try:
            _insert_image_to_cell(cell, img_path, max_width_inches=img_width_inches)
            matched += 1
            print(f"  [OK]  {cam_name}  → Row {row_idx}")
        except Exception as e:
            print(f"  [ERR] {cam_name} 插图失败: {e}")
            unmatched.append(cam_name)

    # 确保输出目录存在
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"[完成] 台账已保存: {output_path}")
    print(f"  插入: {matched}  未匹配: {len(unmatched)}  跳过: {skipped}")

    return {"matched": matched, "unmatched": unmatched, "skipped": skipped}


# ─── CLI 入口 ─────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="将巡检截图贴入 Word 巡查台账")
    parser.add_argument("--docx", required=True, help="输入 Word 模版路径")
    parser.add_argument("--results", default=None, help="巡检 JSON 结果文件路径")
    parser.add_argument("--output", required=True, help="输出 Word 文件路径")
    parser.add_argument("--date", default=None, help="巡查日期（如 2026.6.29），默认今天")
    parser.add_argument("--img-width", type=float, default=1.8, help="图片宽度（英寸），默认 1.8")
    parser.add_argument("--include-offline", action="store_true", help="包含离线摄像头行（不插图但不跳过）")
    args = parser.parse_args()

    if not args.results:
        print("ERROR: 必须通过 --results 指定巡检 JSON 结果文件")
        sys.exit(1)

    if not os.path.exists(args.results):
        print(f"ERROR: 结果文件不存在: {args.results}")
        sys.exit(1)

    with open(args.results, "r", encoding="utf-8") as f:
        results = json.load(f)

    print(f"[update_docx] 载入结果: {len(results)} 条")
    print(f"[update_docx] 模版: {args.docx}")
    print(f"[update_docx] 输出: {args.output}")

    stat = update_docx(
        docx_path=args.docx,
        results=results,
        output_path=args.output,
        today_str=args.date,
        img_width_inches=args.img_width,
        skip_offline=not args.include_offline,
    )

    sys.exit(0 if not stat["unmatched"] else 2)


if __name__ == "__main__":
    main()
